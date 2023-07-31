from django.http import Http404
from django.shortcuts import get_object_or_404
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework.decorators import action
from.models import Article, User
from.serializers import ArticleSerializer, UserSerializer
from django.http import FileResponse
from docx import Document
from io import BytesIO
from django.views.generic import View




class UserListView(APIView):
    
    def get(self,request, formart=None):
        user = User.objects.all()
        serializer = UserSerializer(user, many=True)
        return Response(serializer.data)
    
class UserDetailsView(APIView):
    def get_object(self, pk):
        try:
            return User.objects.get(pk=pk)
        except User.DoesNotExist:
            raise Http404     
        
    def get(self, request, pk, format=None):
        user = self.get_object(pk)
        serializer = UserSerializer(user)
        return Response(serializer.data)    
 

class ArticleListView(APIView):
    def get(self, request, format=None):
        articles = Article.objects.all()
        serializer = ArticleSerializer(articles, many=True)
        return Response(serializer.data)
    

class ArticleDetailsView(APIView): 
    def get_object(self, pk): 
        try:
            return Article.objects.get(pk=pk)
        except Article.DoesNotExist:
            raise Http404
              
    def get(self, request, pk, format=None):
        article = self.get_object(pk)
        serializer = ArticleSerializer(article)
        return Response(serializer.data)


class DownloadArticleView(APIView):
# Here we are getting the particular article using the article id 
    def get(self, request, article_id):
        
        article_content = request.GET.get('content', '')
        # getting the content of the article like the article body 

        # Remove HTML tags from the content using a regular expression
        import re
        plain_text_content = re.sub('<[^>]+>', '', article_content)

        # Generate the Word document
        document = Document()
        #get the article object from the database 
        article = get_object_or_404(Article, id=article_id)  # Import the get_object_or_404 function
        document.add_heading(article.title, level=1)  # Add the title as heading
        document.add_paragraph(plain_text_content)  # Add the plain text content as a paragraph

        # Save the Word document to a BytesIO object (in-memory file)
        output = BytesIO()
        document.save(output)
        output.seek(0)

        # Create a FileResponse and send the Word document as a downloadable file
        response = FileResponse(output, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename=article_{article_id}.docx'
        return response
